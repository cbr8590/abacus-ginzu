"""
Extract text from company diligence documents (PDF, Word, txt).
"""

from pathlib import Path
from typing import Union

try:
    import ginzu_debug
except ImportError:
    ginzu_debug = None

def _log(msg: str, level: str = "info"):
    if ginzu_debug:
        ginzu_debug.log(msg, level)


def extract_text(file_path: Union[str, Path]) -> str:
    """
    Extract text from a document. Supports PDF, Word (.docx), and plain text.
    
    Returns:
        Extracted text content, or empty string if unsupported/failed.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")
    
    suffix = path.suffix.lower()
    _log(f"extract_text: {path.name} (format={suffix})")
    
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8", errors="replace")
        _log(f"Extracted {len(text)} chars from txt")
        return text
    
    if suffix == ".pdf":
        text = _extract_pdf(path)
        _log(f"Extracted {len(text)} chars from PDF")
        return text
    
    if suffix == ".docx":
        text = _extract_docx(path)
        _log(f"Extracted {len(text)} chars from Word")
        return text
    
    raise ValueError(
        f"Unsupported format: {suffix}. Use .pdf, .docx, or .txt"
    )


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError(
            "pypdf required for PDF support. Run: pip install pypdf"
        )
    
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    """Extract text from Word document (paragraphs + tables)."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx required for Word support. Run: pip install python-docx"
        )
    
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)
